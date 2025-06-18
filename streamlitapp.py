import streamlit as st
from utils.pdf_loader import extract_resume_text
from utils.text_splitter import split_text
from utils.embedder import embed_texts, embed_query
from utils.pinecone_store import upsert_embeddings, query_pinecone
from openai import OpenAI
from config import OPENAI_API_KEY
import os
import csv
from io import StringIO, BytesIO
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="Resume Q&A", layout="wide")
st.title("📄 Resume Q&A with RAG")

# Initialize OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)

# Session state
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "context" not in st.session_state:
    st.session_state.context = ""
if "chunks" not in st.session_state:
    st.session_state.chunks = []

# Sidebar
with st.sidebar:
    st.markdown("### 📎 Upload Resumes")
    uploaded_files = st.file_uploader("Upload PDFs", type="pdf", accept_multiple_files=True)

    st.markdown("---")
    st.markdown("### 🧾 Export Chat")
    export_format = st.radio("Choose format:", [".txt", ".csv", ".docx"])

    if st.button("📥 Download Chat History"):
        if not st.session_state.chat_history:
            st.warning("No chat history to export.")
        elif export_format == ".txt":
            chat_text = "\n\n".join([
                f"User: {msg['content']}" if msg['role'] == 'user' else f"AI: {msg['content']}"
                for msg in st.session_state.chat_history
            ])
            st.download_button("Download TXT", data=chat_text, file_name="resume_chat.txt", mime="text/plain")

        elif export_format == ".csv":
            csv_buf = StringIO()
            writer = csv.writer(csv_buf)
            writer.writerow(["Role", "Message"])
            for msg in st.session_state.chat_history:
                writer.writerow([msg["role"], msg["content"]])
            st.download_button("Download CSV", data=csv_buf.getvalue(), file_name="resume_chat.csv", mime="text/csv")

        elif export_format == ".docx":
            doc = Document()
            doc.add_heading("Resume Chat History", level=1)
            for msg in st.session_state.chat_history:
                p = doc.add_paragraph()
                p.add_run("User: " if msg["role"] == "user" else "AI: ").bold = True
                p.add_run(msg["content"]).font.size = Pt(11)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("Download DOCX", data=buffer, file_name="resume_chat.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.markdown("---")
    st.markdown("### ✨ Career Tools")
    if st.button("📝 Generate Cover Letter"):
        st.session_state.show_cover_form = True
    if st.button("📋 Match Resume to Job Description"):
        st.session_state.show_match_form = True

# Upload + Embed
if uploaded_files:
    all_chunks = []
    for file in uploaded_files:
        resume_text = extract_resume_text(file)
        st.session_state.context += resume_text + "\n"
        chunks = split_text(resume_text)
        all_chunks.extend(chunks)

    st.session_state.chunks = all_chunks
    st.success(f"✅ Loaded {len(uploaded_files)} resumes with {len(all_chunks)} chunks.")

    embeddings = embed_texts(all_chunks)
    upsert_embeddings(all_chunks, embeddings)

# Resume Q&A
st.markdown("### 💬 Ask Questions about Your Resume")
user_query = st.text_input("You:")
if user_query:
    if not st.session_state.chunks:
        st.warning("Upload resumes first.")
    else:
        st.session_state.chat_history.append({"role": "user", "content": user_query})
        query_vector = embed_query(user_query)
        result = query_pinecone(query_vector)
        context = "\n".join([m["metadata"]["text"] for m in result["matches"]])

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful assistant answering resume questions."},
                {"role": "user", "content": f"Use the resume below to answer:\n\n{context}\n\nQ: {user_query}"}
            ]
        )
        answer = response.choices[0].message.content
        st.session_state.chat_history.append({"role": "assistant", "content": answer})

# Chat UI
for msg in st.session_state.chat_history:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Cover Letter Tool
if st.session_state.get("show_cover_form"):
    if not st.session_state.context:
        st.warning("Upload resume first.")
        st.stop()

    st.subheader("📝 Cover Letter Generator")
    job_title = st.text_input("Job Title")
    company = st.text_input("Company Name")
    tone = st.selectbox("Tone", ["Formal", "Friendly", "Enthusiastic"])

    if st.button("Generate Cover Letter"):
        prompt = f"""Write a {tone.lower()} cover letter for the role of {job_title} at {company}.
Use the following resume content:
{st.session_state.context}
"""
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        letter = response.choices[0].message.content
        st.text_area("📄 Cover Letter", value=letter, height=300)

# Job Match Tool
if st.session_state.get("show_match_form"):
    if not st.session_state.context:
        st.warning("Upload resume first.")
        st.stop()

    st.subheader("📋 Job Description Matcher")
    jd = st.text_area("Paste Job Description Here")

    if st.button("Check Match"):
        jd_vector = embed_query(jd)
        result = query_pinecone(jd_vector, top_k=5)
        score = round(sum(m["score"] for m in result["matches"]) / len(result["matches"]) * 100, 2)

        st.success(f"✅ Resume Match Score: {score}%")
        st.markdown("### 🔍 Matching Resume Snippets")
        for m in result["matches"]:
            st.markdown(f"- {m['metadata']['text']} (Score: {round(m['score']*100, 2)}%)")
